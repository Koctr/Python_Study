# -*-coding:utf-8-*-
__author__ = "Koctr"

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

doc = Document()
styles = doc.styles

for style in styles:
    if style.type == WD_STYLE_TYPE.TABLE:
        doc.add_paragraph("表格样式：%s" % style.name)
        table = doc.add_table(3, 3, style=style)
        doc.add_paragraph("\n")

doc.save("C:/000/表格样式.docx")
