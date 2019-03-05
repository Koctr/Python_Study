# -*- coding: utf-8 -*-
# author = 'K0ctr'

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

from win32com.client import Dispatch, constants, gencache
import os
import time

price = input("请输入今日价格：")
company_list = ['客户1', '客户2', '客户3', '客户4', '客户5', '客户6', '客户7', '客户8', '客户9', '客户10']

today = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y='年', m='月', d='月')

for company in company_list:
    document = Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    document.add_picture('D:/banner.jpg', width=Inches(6))

    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('关于下达%s产品价格的通知' % today)
    run1.font.name = u'微软雅黑'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    run1.font.size = Pt(21)
    run1.font.bold = True
    p1.space_after = Pt(5)
    p1.space_before = Pt(5)

    p2 = document.add_paragraph()
    run2 = p2.add_run(company + ': ')
    run2.font.name = u'仿宋_GB2312'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)
    run2.font.bold = True

    p3 = document.add_paragraph()
    run3 = p3.add_run('   根据公司安排，为提供优质客户服务，我单位现将价格通知如下。')
    run3.font.name = u'仿宋_GB2312'
    run3.font.size = Pt(16)
    run3.font.bold = True

    table = document.add_table(rows=3, cols=3, style='Table Grid')

    table.cell(0, 0).merge(table.cell(0, 2))
    table_run1 = table.cell(0, 0).paragraphs[0].add_run('XX产品报价表')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_run1.font.name = u'隶书'
    table_run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'隶书')

    table.cell(1, 0).text = '日期'
    table.cell(1, 1).text = '价格'
    table.cell(1, 2).text = '备注'
    table.cell(2, 0).text = today
    table.cell(2, 1).text = str(price)
    table.cell(2, 2).text = ''

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('（联系人：小K    电话：18888888888')
    run4.font.name = u'仿宋_GB2312'
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)
    run4.font.bold = True

    document.add_page_break()
    p5 = document.add_paragraph()
    run5 = p5.add_run('此处是广告')

    if os.path.exists('D:/%s-价格通知.docx' % company):
        os.remove('D:/%s-价格通知.docx' % company)
    document.save("D:/%s-价格通知.docx" % company)

    docx_path = 'D:/%s-价格通知.docx' % company
    pdf_path = 'D:/%s-价格通知.pdf' % company

    gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)
    wd = Dispatch("Word.Application")

    docx_file = wd.Documents.Open(docx_path, ReadOnly=1)
    docx_file.ExportAsFixedFormat(pdf_path, constants.wdExportFormatPDF, Item=constants.wdExportDocumentWithMarkup,
                                  CreateBookmarks=constants.wdExportCreateHeadingBookmarks)

    wd.Quit(constants.wdDoNotSaveChanges)
    time.sleep(5)
    os.remove("D:/%s-价格通知.docx" % company)
