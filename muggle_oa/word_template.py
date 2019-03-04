# -*- coding: utf-8 -*-
# author = 'K0ctr'

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import random


def add_context(document, context):
    document.styles["Normal"].font.name = u'仿宋GB2312'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋GB2312')

    styles = document.styles
    paragraph_styles = [style for style in styles if style.type == WD_STYLE_TYPE.PARAGRAPH]

    p = document.add_paragraph()
    p.style = paragraph_styles[random.randint(0, len(paragraph_styles) - 1)]
    r = p.add_run(str(context))


def change_text(document, old_text, new_text):
    all_paragraphs = document.paragraphs
    for p in all_paragraphs:
        for run in p.runs:
            run_text = run.text.replace(old_text, new_text)
            run.text = run_text
    all_tables = document.tables
    for t in all_tables:
        for row in t.rows:
            for cell in row.cells:
                cell_text = cell.text.replace(old_text, new_text)
                cell.text = cell_text


doc = Document()
add_context(doc, '汉皇重色思倾国，御宇多年求不得。杨家有女初长成，养在深闺人未识。')
add_context(doc, '天生丽质难自弃，一朝选在君王侧。回眸一笑百媚生，六宫粉黛无颜色。')
add_context(doc, '春寒赐浴华清池，温泉水滑洗凝脂。侍儿扶起娇无力，始是新承恩泽时。')
add_context(doc, '云鬓花颜金步摇，芙蓉帐暖度春宵。春宵苦短日高起，从此君王不早朝。')
add_context(doc, '承欢侍宴无闲暇，春从春游夜专夜。后宫佳丽三千人，三千宠爱在一身。')
add_context(doc, '金屋妆成娇侍夜，玉楼宴罢醉和春。姊妹弟兄皆列土，可怜光彩生门户。')
add_context(doc, '遂令天下父母心，不重生男重生女。骊宫高处入青云，仙乐风飘处处闻。')
add_context(doc, '缓歌慢舞凝丝竹，尽日君王看不足。渔阳鼙鼓动地来，惊破霓裳羽衣曲。')
add_context(doc, '九重城阙烟尘生，千乘万骑西南行。翠华摇摇行复止，西出都门百余里。')
add_context(doc, '六军不发无奈何，宛转蛾眉马前死。花钿委地无人收，翠翘金雀玉搔头。')
add_context(doc, '君王掩面救不得，回看血泪相和流。')

change_text(doc, '女', 'girl')
doc.save("D:/长恨歌_替换.docx")
