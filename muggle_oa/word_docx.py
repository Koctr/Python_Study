# -*-coding:utf-8-*-
# Author: Koctr

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn


def set_document(document, content):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.space_before = Pt(5)
    paragraph.space_after = Pt(5)
    run = paragraph.add_run(content)
    run.font.name = u'黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.size = Pt(16)
    run.font.bold = True


document = Document()
content = "这是1个段落。"
for i in range(10):
    set_document(document, content)
document.save("C:/000/paragraph.docx")
