# -*- coding: utf-8 -*-
# author = 'K0ctr'

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
import zipfile

doc = zipfile.ZipFile("D:/长恨歌3.docx")
xml = doc.read("word/document.xml").decode(encoding="utf-8")
xml_list = xml.split("<w:t>")
text_list = []
for text in xml_list:
    if text.find("</w:t>") != -1:
        text_list.append(text[:text.find("</w:t>")])
    else:
        pass

new_doc = Document()
new_doc.styles["Normal"].font.name = u'微软雅黑'
new_doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
paragraph = new_doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for text in text_list:
    run = paragraph.add_run(text)
    run.font.size = Pt(12)

new_doc.save("D:/长恨歌.docx")
