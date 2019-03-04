# -*- coding: utf-8 -*-
# author = 'K0ctr'

from docx import Document
import xlrd


def change_text(document, old_text, new_text):
    all_paragraph = document.paragraphs
    for p in all_paragraph:
        for r in p.runs:
            run_text = r.text.replace(old_text, new_text)
            r.text = run_text

    all_tables = document.tables
    for t in all_tables:
        for trow in t.rows:
            for cell in trow.cells:
                cell_text = cell.text.replace(old_text, new_text)
                cell.text = cell_text


doc = Document("D:/长恨歌3.docx")

excel = xlrd.open_workbook("D:/change.xlsx")
sheet = excel.sheet_by_index(0)
for row in range(1, sheet.nrows):
    for col in range(0, sheet.ncols):
        change_text(doc, str(sheet.cell_value(0, col)), str(sheet.cell_value(row, col)))

doc.save("D:/长恨歌替换.docx")
